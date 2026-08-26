from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Helper functions ────────────────────────────────────────────────────────

def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade = OxmlElement('w:shd')
        shade.set(qn('w:fill'), '1F3864')
        shade.set(qn('w:color'), 'FFFFFF')
        hdr[i]._tc.get_or_add_tcPr().append(shade)
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            row[ci].text = val
            if ri % 2 == 1:
                shade = OxmlElement('w:shd')
                shade.set(qn('w:fill'), 'E9ECF5')
                row[ci]._tc.get_or_add_tcPr().append(shade)
    doc.add_paragraph()
    return table

def add_divider(doc):
    p = doc.add_paragraph('─' * 80)
    p.runs[0].font.color.rgb = RGBColor(150, 150, 150)

def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.left_indent = Inches(0.3)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')
    p._p.get_or_add_pPr().append(shd)
    return p

# ─── COVER PAGE ──────────────────────────────────────────────────────────────

doc.add_paragraph()
title = doc.add_heading('AI Risk Manager', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0, 112, 192)

subtitle = doc.add_paragraph('Chargeback Evidence Responder — Complete Study Guide')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(16)

meta = doc.add_paragraph('Razorpay Internship Hackathon  |  August 2026')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

# ─── SECTION 1: BIG PICTURE ──────────────────────────────────────────────────

set_heading(doc, 'SECTION 1: The Big Picture — What Problem Are We Solving?', 1, (0, 112, 192))
add_paragraph(doc,
    'Razorpay is a payment gateway that sits between merchants (online stores) and banks. '
    'Every day, three types of losses silently drain merchant revenue: Fraud, Returns, and Chargebacks. '
    'Our project — the Chargeback Evidence Responder — tackles chargebacks specifically, '
    'because 60–80% of them are "friendly fraud" that merchants can WIN if they respond correctly.',
    size=11)

doc.add_paragraph()
add_table(doc,
    ['Problem', 'What Happens', 'Who Loses Money'],
    [
        ('Fraud', 'Stolen card / fake identity used to buy products', 'Merchant + Payment Gateway'),
        ('Returns', 'Customer returns wrong/fake item and keeps the original', 'Merchant'),
        ('Chargebacks', 'Customer tells bank "I didn\'t authorize this" → bank reverses payment', 'Merchant pays penalty + loses product'),
    ]
)

# ─── SECTION 2: CHARGEBACK LIFECYCLE ─────────────────────────────────────────

set_heading(doc, 'SECTION 2: The Full Life of a Chargeback', 1, (0, 112, 192))

set_heading(doc, '2.1 Cast of Characters', 2)
add_table(doc,
    ['Player', 'Role'],
    [
        ('🛍️ Customer', 'Buys something online'),
        ('🏪 Merchant', 'Sells the product (e.g. a Myntra seller)'),
        ('💳 Card Bank', 'Customer\'s bank (e.g. HDFC, SBI)'),
        ('🏦 Acquirer', 'Merchant\'s bank — processes payments for the merchant'),
        ('⚡ Razorpay', 'Payment gateway sitting in the middle'),
    ]
)

set_heading(doc, '2.2 Step-by-Step Chargeback Timeline', 2)
steps = [
    'Day 1:   Customer buys ₹10,000 phone via Razorpay ✅',
    'Day 2:   Money flows: Customer\'s Bank → Razorpay → Merchant ✅',
    'Day 45:  Customer calls HDFC: "I never made this purchase"',
    'Day 46:  HDFC files a CHARGEBACK against Razorpay',
    'Day 47:  Razorpay alerts merchant: "You have 7 days to respond with evidence"',
    'Day 54:  If merchant fails to respond → HDFC takes back ₹10,000 + ₹1,500 fee',
    '         If merchant responds well → Case reviewed, merchant can WIN',
]
for s in steps:
    add_bullet(doc, s)

set_heading(doc, '2.3 Why Do Chargebacks Happen?', 2)
add_table(doc,
    ['Reason', 'Explanation', 'Actually Fraud?'],
    [
        ('True Fraud', 'Stolen card — customer genuinely didn\'t transact', 'YES'),
        ('Friendly Fraud', 'Customer bought it, liked it, then disputed it', 'NO (abuse)'),
        ('Non-receipt', 'Customer claims product never arrived', 'MAYBE'),
        ('Not as Described', 'Product different from what was sold', 'MAYBE'),
        ('Duplicate Processing', 'Charged twice by mistake', 'NO'),
        ('Subscription Cancelled', 'Recurring charge after cancellation', 'DEPENDS'),
    ]
)
add_paragraph(doc,
    '⚠️  KEY INSIGHT: Studies show 60–80% of chargebacks are "friendly fraud" — '
    'the customer actually did make the purchase. Merchants can WIN most cases '
    'if they fight them with the right evidence. This is exactly what our AI does.',
    bold=True)

# ─── SECTION 3: EVIDENCE ─────────────────────────────────────────────────────

set_heading(doc, 'SECTION 3: What Evidence Wins a Chargeback?', 1, (0, 112, 192))

add_paragraph(doc,
    'When a merchant responds to a chargeback, they submit a Rebuttal Letter + Evidence Package. '
    'Evidence is tiered by strength:', size=11)

set_heading(doc, '3.1 Tier 1 — Strongest Evidence (Always Include)', 2)
add_table(doc,
    ['Evidence', 'What It Proves'],
    [
        ('Transaction timestamp + IP address', 'Purchase came from a real device at a real location'),
        ('Device fingerprint', 'Unique identifier of the device used to buy'),
        ('Billing address match (AVS)', 'Address provided matches the card\'s billing address'),
        ('CVV / 3D Secure OTP verified', 'Customer passed all security checks — OTP went to their phone'),
        ('Order confirmation email sent', 'They had access to the email on the account'),
    ]
)

set_heading(doc, '3.2 Tier 2 — Strong Evidence', 2)
add_table(doc,
    ['Evidence', 'What It Proves'],
    [
        ('Delivery confirmation + courier tracking', 'Physical product delivered to their address'),
        ('Signed delivery receipt', 'Someone at the address accepted the package'),
        ('Previous successful transactions', 'This customer has purchased from us before'),
        ('Login activity logs', 'Customer logged in, browsed, then purchased'),
        ('Customer service interaction', 'They contacted support after the purchase'),
    ]
)

set_heading(doc, '3.3 Tier 3 — Supporting Evidence', 2)
add_table(doc,
    ['Evidence', 'What It Proves'],
    [
        ('Post-delivery product usage', 'They used/activated the product (e.g. app login)'),
        ('Social media / review', 'Customer shared or reviewed the product'),
        ('Refund policy acknowledgment', 'Customer agreed to return terms at checkout'),
    ]
)

set_heading(doc, '3.4 The "Smoking Gun" Pieces', 2)
add_paragraph(doc, 'Some evidence almost single-handedly wins the case:', bold=True)
add_bullet(doc, '3D Secure OTP: Customer claims "Not my transaction" — but the OTP went to their registered mobile. How did someone else get it?', bold_prefix='3DS OTP → ')
add_bullet(doc, 'Signed delivery proof: Customer claims "I never received it" — but someone signed for it at their address.', bold_prefix='Signed Receipt → ')
add_bullet(doc, 'Post-purchase usage: Customer says "I didn\'t authorize this Netflix charge" — but 47 hours of streaming happened after billing.', bold_prefix='Product Usage → ')

# ─── SECTION 4: HOW EVIDENCE LOGICALLY COUNTERS CLAIMS ───────────────────────

set_heading(doc, 'SECTION 4: How Evidence Logically Disproves Customer Claims', 1, (0, 112, 192))

add_paragraph(doc,
    'Banks use a "balance of probability" standard — not "beyond reasonable doubt." '
    'They ask: "Is it MORE likely this was legitimate, or fraud?" '
    'Our job is to make the customer\'s story implausible.', size=11)

set_heading(doc, '4.1 Scenario: Customer says "I never made this purchase"', 2)
add_table(doc,
    ['Evidence We Have', 'Logical Conclusion for the Bank'],
    [
        ('IP Address = customer\'s city', 'Transaction came from WHERE they live'),
        ('Device fingerprint = their regular phone', 'Purchase happened on the device they always use'),
        ('3DS OTP was verified', 'Only they could receive and enter the OTP'),
        ('Billing address matched', 'Whoever bought it knew their home address'),
        ('Logged in 2 hours after purchase', 'They were actively using the account that day'),
        ('Product delivered to their home, signed', 'A real person at their address accepted it'),
        ('Contacted support asking about shipping', 'Why would a fraud victim ask about shipping?'),
    ]
)

set_heading(doc, '4.2 The Bank\'s Final Decision Logic', 2)
add_code_block(doc,
    'Customer Claim:  "I never made this purchase"\n\n'
    'Bank sees:\n'
    '  ✅ OTP entered from customer\'s own phone\n'
    '  ✅ Delivery signed at their home address\n'
    '  ✅ Customer emailed us asking about delivery\n'
    '  ✅ Same device used for 15 previous purchases\n\n'
    'Bank decision: "This claim is IMPLAUSIBLE → Chargeback DENIED → Merchant WINS"'
)

# ─── SECTION 5: AI SYSTEM ARCHITECTURE ───────────────────────────────────────

set_heading(doc, 'SECTION 5: What Our AI System Does — 5-Step Pipeline', 1, (0, 112, 192))

add_paragraph(doc,
    'Our system is a semi-automated chargeback response pipeline. '
    'Think of it as a smart law firm for merchants:', size=11)

steps_ai = [
    ('Step 1: Risk Classifier', 'ML model predicts Win Probability (0.0–1.0). '
        'Score > 0.6 → Fight it. Score 0.3–0.6 → Human review. Score < 0.3 → Accept the loss.'),
    ('Step 2: Evidence Collector', 'Automatically pulls transaction logs, delivery status, '
        'customer activity logs, and past transaction history.'),
    ('Step 3: Evidence Scorer', 'Rates how strong our evidence package is. '
        'Outputs STRONG / MODERATE / WEAK.'),
    ('Step 4: Rebuttal Letter Generator', 'AI generates a professional dispute letter '
        'customized by reason code, ready to submit to the bank as PDF/DOCX.'),
    ('Step 5: Metrics & Reporting', 'Tracks win rate, false positive cost (in ₹), '
        'and total value recovered over time.'),
]
for title_s, desc in steps_ai:
    add_bullet(doc, f' {desc}', bold_prefix=title_s + ':')

# ─── SECTION 6: ML MODEL ─────────────────────────────────────────────────────

set_heading(doc, 'SECTION 6: The ML Model — Win Probability Classifier', 1, (0, 112, 192))

set_heading(doc, '6.1 What We\'re Predicting', 2)
add_paragraph(doc, 'For each incoming chargeback:', size=11)
add_bullet(doc, 'Label 1 = Merchant will WIN this chargeback (if they respond)')
add_bullet(doc, 'Label 0 = Merchant will LOSE even if they respond')

set_heading(doc, '6.2 Features (Model Inputs)', 2)
add_table(doc,
    ['Feature', 'Description'],
    [
        ('transaction_amount', 'High amounts are harder to win'),
        ('days_since_transaction', 'Older transactions are harder to dispute'),
        ('chargeback_reason_code', 'Different codes have different win rates'),
        ('is_3ds_verified', '3D Secure verified — very strong defense'),
        ('avs_match', 'Address Verification match (Y/N)'),
        ('cvv_match', 'CVV matched at time of transaction'),
        ('customer_age_days', 'How long this customer account has existed'),
        ('previous_orders_count', 'First-time buyer vs. loyal customer'),
        ('previous_chargebacks_count', 'Is this customer a serial disputer?'),
        ('has_delivery_confirmation', 'Do we have proof of delivery?'),
        ('has_signed_receipt', 'Signed delivery proof exists'),
        ('has_customer_login_after', 'Did they log in after the purchase?'),
        ('has_support_interaction', 'Did they contact support about the order?'),
        ('evidence_completeness_score', '% of key evidence we have (0.0–1.0)'),
        ('merchant_win_rate', 'This merchant\'s historical win rate'),
        ('reason_code_win_rate', 'Win rate for this specific reason code'),
    ]
)

set_heading(doc, '6.3 Algorithm: XGBoost Classifier', 2)
add_paragraph(doc,
    'We use XGBoost (Extreme Gradient Boosting) — the most powerful algorithm for tabular data. '
    'It handles missing data gracefully, gives feature importance for explainability, '
    'and is the industry standard for payment fraud competitions.', size=11)

# ─── SECTION 7: METRICS ───────────────────────────────────────────────────────

set_heading(doc, 'SECTION 7: Metrics — Precision, Recall & False Positive Cost', 1, (0, 112, 192))

set_heading(doc, '7.1 Confusion Matrix (Example)', 2)
add_table(doc,
    ['', 'Actually WIN', 'Actually LOSE'],
    [
        ('Model says FIGHT', 'True Positive (TP) — Correctly fought & won', 'False Positive (FP) — Fought & LOST (wasted effort)'),
        ('Model says SKIP', 'False Negative (FN) — Missed a winnable case', 'True Negative (TN) — Correctly skipped'),
    ]
)

set_heading(doc, '7.2 Formulas', 2)
add_code_block(doc,
    'Precision = TP / (TP + FP)\n'
    '          = "Of all chargebacks we decided to fight, what % did we WIN?"\n\n'
    'Recall    = TP / (TP + FN)\n'
    '          = "Of all winnable chargebacks, what % did we correctly identify?"\n\n'
    'F1 Score  = 2 × (Precision × Recall) / (Precision + Recall)\n'
    '          = Balanced metric — the one number that summarizes model quality'
)

set_heading(doc, '7.3 The False Positive Cost (Critical for Razorpay)', 2)
add_paragraph(doc,
    '⚠️  Razorpay explicitly requires "honest metrics including false-positive cost." '
    'This is the cost of WRONGLY telling a merchant to fight a chargeback they will lose:', bold=True)
add_bullet(doc, 'Merchant spends time preparing the evidence package')
add_bullet(doc, 'Pays dispute response submission fee')
add_bullet(doc, 'Still loses the original chargeback amount')
add_bullet(doc, 'Pays chargeback penalty fee (~₹1,000–₹2,500 per case)')
add_paragraph(doc,
    'We calculate this cost in ₹ for every wrong prediction and report it in the dashboard. '
    'This is what separates a good submission from a great one.', italic=True)

# ─── SECTION 8: PROJECT STRUCTURE ─────────────────────────────────────────────

set_heading(doc, 'SECTION 8: Project Structure & Build Roadmap', 1, (0, 112, 192))

set_heading(doc, '8.1 Folder Structure', 2)
add_code_block(doc,
    'chargeback-evidence-responder/\n'
    '│\n'
    '├── data/\n'
    '│   ├── generate_dataset.py      ← Generates 50K synthetic chargeback records\n'
    '│   ├── chargebacks_raw.csv\n'
    '│   └── chargebacks_processed.csv\n'
    '│\n'
    '├── model/\n'
    '│   ├── train_model.py           ← Trains XGBoost classifier\n'
    '│   ├── evaluate_model.py        ← Precision, Recall, FP cost calculation\n'
    '│   └── model.pkl                ← Saved trained model\n'
    '│\n'
    '├── evidence/\n'
    '│   ├── evidence_collector.py\n'
    '│   └── evidence_scorer.py\n'
    '│\n'
    '├── responder/\n'
    '│   ├── letter_generator.py      ← AI rebuttal letter generator\n'
    '│   └── templates/               ← Letter templates per reason code\n'
    '│\n'
    '├── api/\n'
    '│   └── main.py                  ← FastAPI endpoint\n'
    '│\n'
    '├── dashboard/\n'
    '│   └── app.py                   ← Streamlit demo dashboard\n'
    '│\n'
    '├── requirements.txt\n'
    '└── README.md'
)

set_heading(doc, '8.2 Build Roadmap', 2)
add_table(doc,
    ['Phase', 'What We Build', 'Time Estimate'],
    [
        ('Phase 1', 'Python environment setup + project scaffold', '30 mins'),
        ('Phase 2', 'Synthetic dataset generation (50K records)', '1 hour'),
        ('Phase 3', 'Feature engineering + Exploratory Data Analysis', '1–2 hours'),
        ('Phase 4', 'XGBoost model training + evaluation', '1–2 hours'),
        ('Phase 5', 'Evidence collector + scorer', '1 hour'),
        ('Phase 6', 'Rebuttal letter generator', '1–2 hours'),
        ('Phase 7', 'Streamlit dashboard', '1–2 hours'),
        ('Phase 8', 'FastAPI endpoint', '30 mins'),
        ('Phase 9', 'Final metrics + README', '1 hour'),
    ]
)

# ─── SECTION 9: TECH STACK ─────────────────────────────────────────────────────

set_heading(doc, 'SECTION 9: Tech Stack', 1, (0, 112, 192))
add_table(doc,
    ['Library', 'What It Does', 'When We Use It'],
    [
        ('pandas', 'Handle data like Excel but in code', 'Load and clean datasets'),
        ('numpy', 'Math operations', 'Feature calculations'),
        ('scikit-learn', 'Classic ML models + evaluation tools', 'Metrics, train/test split'),
        ('xgboost', 'Powerful tree-based model', 'Main win-probability classifier'),
        ('matplotlib / seaborn', 'Charts and graphs', 'Visualizing results'),
        ('fastapi', 'Build web APIs', 'Serving the model as an endpoint'),
        ('streamlit', 'Quick web dashboards', 'Demo UI for judges'),
        ('shap', 'Explainability', 'Showing WHY a chargeback is flagged'),
        ('python-docx', 'Generate Word documents', 'Auto-generating rebuttal letters'),
    ]
)

# ─── SECTION 10: GLOSSARY ─────────────────────────────────────────────────────

set_heading(doc, 'SECTION 10: Glossary — Quick Reference', 1, (0, 112, 192))
add_table(doc,
    ['Term', 'Meaning'],
    [
        ('BFSI', 'Banking, Financial Services, Insurance'),
        ('Chargeback', 'Forced reversal of payment by the customer\'s bank'),
        ('Friendly Fraud', 'Legitimate customer disputes a real transaction to get money back'),
        ('False Positive', 'We flag a chargeback as WINNABLE but we actually LOSE it'),
        ('False Negative', 'We flag a chargeback as SKIP but we could have WON it'),
        ('Precision', '% of chargebacks we decided to fight that we actually won'),
        ('Recall', '% of all winnable chargebacks that we correctly identified'),
        ('F1 Score', 'Balanced metric combining Precision and Recall'),
        ('ROC-AUC', 'Overall model quality (0.5 = random, 1.0 = perfect)'),
        ('XGBoost', 'Extreme Gradient Boosting — our main ML algorithm'),
        ('3DS / 3D Secure', 'OTP-based authentication for online card transactions'),
        ('AVS', 'Address Verification System — checks billing address'),
        ('CVV', '3-digit security code on the back of a card'),
        ('Feature Engineering', 'Creating useful input variables from raw transaction data'),
        ('Win Probability', 'ML model output: probability (0–1) that we\'ll win this chargeback'),
        ('Evidence Completeness Score', '% of key evidence pieces we have for a transaction'),
        ('Rebuttal Letter', 'Formal document submitted to bank to contest the chargeback'),
        ('Entity Resolution', 'Figuring out if two accounts belong to the same person'),
        ('KYC', 'Know Your Customer — identity verification process'),
        ('Velocity', 'How fast something is happening (e.g. 10 transactions in 5 minutes)'),
    ]
)

# ─── FINAL NOTE ───────────────────────────────────────────────────────────────

add_divider(doc)
set_heading(doc, 'Key Takeaway', 2)
add_paragraph(doc,
    'Most merchants lose chargebacks not because they don\'t have a case, but because they '
    'don\'t know how to present their evidence and miss the deadline. '
    'Our AI solves both problems: it identifies which chargebacks are worth fighting '
    'AND automatically builds the evidence package — all within minutes of the chargeback being filed.',
    bold=True, size=12)

add_paragraph(doc,
    'This is a real, measurable, defense-only business problem with clear ROI. Perfect for Razorpay.',
    italic=True)

add_divider(doc)
p = doc.add_paragraph('Razorpay AI Risk Manager Hackathon  |  August 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.color.rgb = RGBColor(130, 130, 130)

# ─── SAVE ─────────────────────────────────────────────────────────────────────

output_path = r'd:\ochrestra\razor hacka\AI_Risk_Manager_Study_Guide.docx'
doc.save(output_path)
print(f'[OK] Document saved to: {output_path}')
