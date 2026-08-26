"""Rebuttal letter generator.

Selects the reason-code template, fills every placeholder with live case
data, builds a numbered evidence schedule from the collected bundle, injects
a deadline urgency notice when the response window is closing, and renders
the finished letter both as text and as a Word document (in-memory bytes).
"""

from __future__ import annotations

import sys
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any

import yaml

PACKAGE_DIR = Path(__file__).resolve().parent.parent
if str(PACKAGE_DIR) not in sys.path:
    sys.path.insert(0, str(PACKAGE_DIR))

DEFAULT_TEMPLATES_DIR = PACKAGE_DIR / "responder" / "templates"
DEFAULT_CONFIG_PATH = PACKAGE_DIR / "config.yaml"

REASON_TO_TEMPLATE = {
    "CB001": "unauthorized.txt",
    "CB002": "non_receipt.txt",
    "CB003": "not_as_described.txt",
    "CB004": "friendly_fraud.txt",
}
REASON_DESCRIPTIONS = {
    "CB001": "Transaction Not Authorised",
    "CB002": "Goods/Services Not Received",
    "CB003": "Not as Described",
    "CB004": "General Dispute Claim",
}

EVIDENCE_LABELS = {
    "is_3ds_verified": "3-D Secure Authentication",
    "has_delivery_confirmation": "Courier Delivery Confirmation",
    "avs_match": "Address Verification Service (AVS) Match",
    "cvv_match": "Card Security Code (CVV) Verification",
    "has_signed_receipt": "Signed Proof of Delivery",
    "has_login_after_purchase": "Post-Purchase Account Login",
    "has_support_interaction": "Customer Support Interaction",
    "order_confirmation_sent": "Order Confirmation Email",
    "refund_policy_acknowledged": "Refund Policy Acknowledgement",
}

URGENT_WINDOW_DAYS = 3


@dataclass
class LetterResult:
    letter_text: str
    word_doc_bytes: bytes
    template_used: str
    missing_evidence_warnings: list[str] = field(default_factory=list)
    urgent: bool = False


def _british_date(value: Any) -> str:
    if isinstance(value, str):
        value = datetime.fromisoformat(value).date()
    elif isinstance(value, datetime):
        value = value.date()
    return value.strftime("%d %B %Y")


class RebuttalLetterGenerator:
    """Builds formal dispute rebuttal letters from templates and live data."""

    def __init__(
        self,
        templates_dir: str | Path = DEFAULT_TEMPLATES_DIR,
        config_path: str | Path = DEFAULT_CONFIG_PATH,
    ) -> None:
        self.templates_dir = Path(templates_dir)
        with open(config_path, "r", encoding="utf-8") as fh:
            self.config = yaml.safe_load(fh)
        self.templates: dict[str, str] = {}
        for code, filename in REASON_TO_TEMPLATE.items():
            path = self.templates_dir / filename
            self.templates[filename] = path.read_text(encoding="utf-8")

    # ------------------------------------------------------------ generation

    def generate(
        self,
        chargeback_data: dict[str, Any],
        evidence_bundle: dict[str, Any],
        prediction_result: dict[str, Any],
    ) -> LetterResult:
        reason_code = str(chargeback_data.get("reason_code", "CB004"))
        template_name = REASON_TO_TEMPLATE.get(reason_code, "friendly_fraud.txt")
        template = self.templates[template_name]

        warnings = self._build_missing_warnings(evidence_bundle)
        evidence_section = self._format_evidence_schedule(evidence_bundle)

        win_prob_pct = int(round(float(prediction_result["win_probability"]) * 100))
        strength = str(prediction_result.get("evidence_strength", "MODERATE"))

        replacements = {
            "{{CHARGEBACK_ID}}": str(chargeback_data.get("chargeback_id", "UNKNOWN")),
            "{{REASON_CODE}}": reason_code,
            "{{TRANSACTION_DATE}}": _british_date(chargeback_data["transaction_date"]),
            "{{AMOUNT_INR}}": f"{float(chargeback_data['amount_inr']):,.2f}",
            "{{MERCHANT_NAME}}": str(chargeback_data.get("merchant_name", "The Merchant")),
            "{{CUSTOMER_NAME}}": str(chargeback_data.get("customer_name", "The Cardholder")),
            "{{RESPONSE_DATE}}": datetime.now().strftime("%d %B %Y"),
            "{{EVIDENCE_LIST}}": evidence_section,
            "{{WIN_PROBABILITY_PCT}}": str(win_prob_pct),
            "{{EVIDENCE_STRENGTH}}": strength,
        }
        letter = template
        for slot, value in replacements.items():
            letter = letter.replace(slot, value)

        deadline = chargeback_data.get("deadline_date")
        urgent = False
        if deadline is not None:
            if isinstance(deadline, str):
                deadline_dt = datetime.fromisoformat(deadline).date()
            elif isinstance(deadline, datetime):
                deadline_dt = deadline.date()
            else:
                deadline_dt = deadline
            days_left = (deadline_dt - datetime.now().date()).days
            if days_left <= URGENT_WINDOW_DAYS:
                urgent = True
                urgency_block = (
                    "*** URGENT — TIME SENSITIVE ***\n"
                    f"This representment must reach the acquiring bank by "
                    f"{_british_date(deadline_dt)} "
                    f"(only {max(days_left, 0)} day(s) remaining). "
                    "Kindly accord this matter priority processing."
                )
                marker = "Dear Sir or Madam,"
                letter = letter.replace(
                    marker, f"{urgency_block}\n\n{marker}", 1
                )

        stripped = letter.rstrip()
        sep = "-" * 66
        prefix = "" if stripped.endswith(sep) else "\n" + sep + "\n"
        footer = (
            f"{prefix}"
            f"AI Confidence Score: {win_prob_pct}% — Recommended Action: "
            f"{prediction_result.get('recommendation', 'REVIEW')}\n"
            "Generated by Chargeback Evidence Responder · Human review advised\n"
        )
        letter = stripped + "\n" + footer

        return LetterResult(
            letter_text=letter,
            word_doc_bytes=self._build_docx(letter),
            template_used=template_name,
            missing_evidence_warnings=warnings,
            urgent=urgent,
        )

    # ----------------------------------------------------------- components

    def _format_evidence_schedule(self, bundle: dict[str, Any]) -> str:
        proofs: list[str] = []
        n = 0

        def add(label: str, proves: str) -> None:
            nonlocal n
            n += 1
            proofs.append(f"{n}. {label}: {proves}")

        if bundle.get("is_3ds_verified"):
            add(
                EVIDENCE_LABELS["is_3ds_verified"],
                "the transaction was authenticated via a one-time password sent to the "
                "customer's registered mobile number; only the account holder could have "
                "completed this step.",
            )
        if bundle.get("has_delivery_confirmation"):
            meta = bundle.get("delivery_metadata", {})
            add(
                EVIDENCE_LABELS["has_delivery_confirmation"],
                f"{meta.get('courier', 'the carrier')} recorded delivery under tracking "
                f"number {meta.get('tracking_number', 'N/A')} on "
                f"{meta.get('delivered_date', 'record')} at the cardholder's nominated address.",
            )
        if bundle.get("avs_match"):
            add(
                EVIDENCE_LABELS["avs_match"],
                "the billing address supplied at checkout matched the issuing bank's "
                "records for the card.",
            )
        if bundle.get("cvv_match"):
            add(
                EVIDENCE_LABELS["cvv_match"],
                "the card security code quoted at payment matched the value held by the "
                "issuer, indicating physical possession of the card details.",
            )
        if bundle.get("has_signed_receipt"):
            meta = bundle.get("signature_metadata", {})
            add(
                EVIDENCE_LABELS["has_signed_receipt"],
                f"a signed acknowledgement of receipt exists for this consignment "
                f"(signed: {meta.get('signer_name', 'recipient')}).",
            )
        if bundle.get("has_login_after_purchase"):
            meta = bundle.get("login_metadata", {})
            add(
                EVIDENCE_LABELS["has_login_after_purchase"],
                f"the customer's account was accessed after the purchase "
                f"({meta.get('login_timestamp', 'recorded time')}, IP "
                f"{meta.get('ip_address', 'recorded')}), consistent with genuine engagement.",
            )
        if bundle.get("has_support_interaction"):
            meta = bundle.get("support_metadata", {})
            add(
                EVIDENCE_LABELS["has_support_interaction"],
                f"the customer raised support ticket {meta.get('ticket_id', 'N/A')} "
                f"(\"{meta.get('subject', 'enquiry')}\", {meta.get('created_date', 'dated')}) — "
                "voluntary contact inconsistent with third-party misuse.",
            )
        if bundle.get("order_confirmation_sent"):
            meta = bundle.get("confirmation_metadata", {})
            add(
                EVIDENCE_LABELS["order_confirmation_sent"],
                f"a confirmation email was dispatched to {meta.get('email', 'the customer')} "
                f"at {meta.get('sent_timestamp', 'the time of order')}; no delivery failure "
                "or objection was ever returned.",
            )
        if bundle.get("refund_policy_acknowledged"):
            add(
                EVIDENCE_LABELS["refund_policy_acknowledged"],
                "the customer expressly accepted the published refund and returns policy "
                "at checkout before payment was authorised.",
            )

        if not proofs:
            return "(No documentary evidence is currently available for this transaction.)"
        return "\n".join(proofs)

    def _build_missing_warnings(self, bundle: dict[str, Any]) -> list[str]:
        weights = self.config["evidence_weights"]
        tier1 = ["is_3ds_verified", "has_delivery_confirmation", "avs_match", "cvv_match"]
        warnings: list[str] = []
        for f in tier1:
            if not bundle.get(f):
                warnings.append(
                    f"CRITICAL: {EVIDENCE_LABELS[f]} is absent — banks weigh this heavily."
                )
        for f in ["has_signed_receipt", "has_login_after_purchase", "has_support_interaction"]:
            if not bundle.get(f):
                warnings.append(f"{EVIDENCE_LABELS[f]} could not be attached.")
        del weights
        return warnings

    @staticmethod
    def _build_docx(text: str) -> bytes:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Courier New"
        style.font.size = Pt(10)

        for line in text.splitlines():
            if line.strip():
                doc.add_paragraph(line.rstrip())
            else:
                doc.add_paragraph("")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # -------------------------------------------------------------- outputs

    def save_docx(self, letter_result: LetterResult, output_path: str | Path) -> Path:
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(letter_result.word_doc_bytes)
        return path

    @staticmethod
    def get_letter_preview(letter_result: LetterResult, max_chars: int = 500) -> str:
        text = letter_result.letter_text
        if len(text) <= max_chars:
            return text
        return text[:max_chars].rstrip() + " … [truncated]"


if __name__ == "__main__":
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8", errors="replace")

    from evidence.collector import EvidenceCollector

    collector = EvidenceCollector(seed=11)
    bundle = collector.collect(
        "TXN-2026-000991",
        chargeback_data={
            "is_3ds_verified": True,
            "avs_match": True,
            "cvv_match": True,
            "has_delivery_confirmation": True,
            "has_signed_receipt": True,
            "customer_name": "Ravi Kumar",
            "days_since_transaction": 21,
        },
    )
    chargeback = {
        "chargeback_id": "CHB-2026-000451",
        "transaction_date": (datetime.now() - timedelta(days=21)).isoformat(),
        "amount_inr": 14299.50,
        "reason_code": "CB001",
        "merchant_name": "TechNova Retail Pvt Ltd",
        "customer_name": "Mr. Ravi Kumar",
        "deadline_date": (datetime.now().date() + timedelta(days=2)).isoformat(),
    }
    prediction = {
        "win_probability": 0.82,
        "recommendation": "FIGHT",
        "evidence_strength": "STRONG",
    }

    gen = RebuttalLetterGenerator()
    result = gen.generate(chargeback, bundle, prediction)

    print("=" * 70)
    print("GENERATED LETTER (template:", result.template_used, "| urgent:", result.urgent, ")")
    print("=" * 70)
    print(result.letter_text)
    print("\nMISSING-EVIDENCE WARNINGS:")
    for w in result.missing_evidence_warnings:
        print("  -", w)
    out = gen.save_docx(result, PACKAGE_DIR / "model" / "artifacts" / "sample_letter_CB001.docx")
    print("\nSaved docx ->", out)
    print("\nPREVIEW:\n" + gen.get_letter_preview(result))
